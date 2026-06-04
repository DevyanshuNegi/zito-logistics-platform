from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "prd-pending"
OUT_FILE = OUT_DIR / "Zito_PRD_Pending_Features_By_User.docx"

ACCENT = "1B3F72"
CYAN = "22D3EE"
DARK = "0B1730"
LIGHT = "EEF4FF"
WARN = "FFF4D6"
GREEN = "EAF7EF"
RED = "FDE8EA"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        set_cell_text(hdr.cells[idx], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr.cells[idx], ACCENT)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            set_cell_text(row.cells[idx], value)
            if idx == 0:
                set_cell_shading(row.cells[idx], LIGHT)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_callout(doc: Document, title: str, text: str, fill: str = LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    run.font.size = Pt(10)
    p.add_run("\n" + text).font.size = Pt(9)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 24, DARK),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 12, ACCENT),
        ("Heading 3", 10.5, DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


ROLE_ROWS = [
    [
        "Customer",
        "Public registration/login, OTP sign-in, booking creation, booking detail, tracking, payments, invoices, support, profile, owned fleet, warehouse discovery and booking.",
        "Customer retention program needs richer loyalty/promo/referral ledger tables; warehouse booking needs provider/data-volume QA; fleet image verification is implemented but needs mobile field testing.",
        "Book courier/PTL/FTL, track live route, pay or view invoice, create support ticket, create owned fleet with camera-only evidence.",
    ],
    [
        "Corporate",
        "Corporate bookings, multi-stop movement planning, contracts, invoices, owned fleet, customer-style warehouse booking, credit exposure enforcement.",
        "Contract exposure and consolidated billing need high-volume regression; approval and credit-limit edge cases need more finance QA; corporate fleet should be tested for owner-scope isolation.",
        "Create booking under contract, verify invoice and credit exposure, manage fleet, verify contract denial when limits are exceeded.",
    ],
    [
        "Driver",
        "Driver dashboard, jobs, dispatch accept/reject endpoints, shift flow, heatmap, earnings, fleet view, booking state transitions, SOS and alerts.",
        "Driver subscription tier UI is pending; mobile driver live-photo identity verification needs end-to-end capture QA; earnings/payout needs payroll-data reconciliation tests.",
        "Start shift, accept/reject assignment, update trip status, use heatmap, inspect earnings, verify subscription UI once added.",
    ],
    [
        "Transporter",
        "Fleet onboarding, mandatory vehicle images/documents, Kenya vehicle brand/model catalog, invoices, marketplace opportunities, owned drivers/vehicle assignment.",
        "Featured listing purchase/extend/cancel UI is pending; transporter dashboard/finance depth needs parity checks; fleet camera-only web capture depends on browser/device behavior.",
        "Create fleet with all required images, link driver, accept marketplace opportunity, verify invoice and vehicle approval workflow.",
    ],
    [
        "Courier Company",
        "Courier dashboard/bookings, dispatch, new movement, scan operations, waybills, owned fleet, invoices, marketplace, multi-load/multi-unload execution.",
        "Mobile and web route parity needs regression; county-to-county scan/waybill chain should be tested with realistic parcel volume; platform fee invoices need finance QA.",
        "Create courier movement, scan parcel checkpoints, generate waybill, review dispatch, manage owned fleet.",
    ],
    [
        "Warehouse Partner",
        "Dashboard, bins, inventory, scan, loss detection, listings, warehouse bookings, marketplace, managed warehouse operations.",
        "Warehouse listing/booking capacity reservation needs stress tests; dock scheduling/GRN/pick-pack advanced PRD flows need deeper implementation audit; mobile warehouse surfaces need parity QA.",
        "Submit listing, admin approves listing, customer books storage, partner manages booking status, scan inventory movement.",
    ],
    [
        "Agent",
        "Agent dashboard, fleet, drivers, marketplace; external supply partner separated from internal agency staff.",
        "Agent commission visibility and opportunity lifecycle need end-to-end QA; driver/fleet onboarding under agent network needs approval and isolation tests.",
        "Register agent, view marketplace, onboard/link drivers, create fleet evidence, verify no agency-staff route access.",
    ],
    [
        "Admin / Super Admin",
        "Users, verification, fleet, marketplace, rate cards, surge, invoices, billing, reconciliation, analytics, audit, fraud, alerts, system health, workforce, agencies, warehouse listings, capacity planning.",
        "Frontend pending for revenue modules: driver subscriptions, featured listings purchase flow, verification expedite button. Multi-region provisioning remains environment-driven. Dedicated promo/referral ledgers may be needed later.",
        "Approve KYC/fleet/listings, publish marketplace opportunity, review fraud, reconcile payments, inspect health, run billing/platform fee flows.",
    ],
    [
        "Agency Staff",
        "Agency login, accounts, operations, support queue, support ticket detail; internal-only access.",
        "Department-specific permissions and agency handoff workflows need full RBAC testing; inter-agency settlement/proof paths need operational test data.",
        "Agency staff login, open support queue, update ticket, verify scoped access to agency data only.",
    ],
]

PRIORITY_ROWS = [
    ["P0", "Revenue frontend gaps", "Driver subscriptions UI, featured listing purchase/extend/cancel UI, verification expedite button are backend-backed but frontend incomplete.", "Admin, Driver, Transporter, Partner", "Build web/mobile entry points and run payment/wallet regression."],
    ["P0", "End-to-end provider QA", "OTP, payment, SMS, map, PDF, scan/offline, and upload provider failures need target-environment validation.", "All users", "Run staging E2E with provider test credentials and document fallback behavior."],
    ["P1", "Mobile parity", "Expo surfaces exist for many roles but need full Android/iPhone parity across customer, driver, transporter, courier, warehouse, internal and agency flows.", "All mobile users", "Device matrix testing, screenshots, and route-by-route smoke pass."],
    ["P1", "Warehouse advanced operations", "PRD describes dock scheduling, GRN, pick-pack-dispatch, cold chain, VAS, cross-docking, inter-warehouse transfer and cycle counts; current repo has strong warehouse base but needs detailed implementation audit for these advanced flows.", "Warehouse Partner, Customer, Admin", "Audit backend models/routes and build missing advanced operational screens if absent."],
    ["P1", "Commercial program ledgers", "Promo code, loyalty points, and referrals currently rely on wallet/audit conventions rather than dedicated program tables.", "Customer, Admin, Finance", "Decide whether to add dedicated promo/loyalty/referral ledger models before scale."],
    ["P1", "Capacity planning precision", "Warehouse capacity uses bin occupancy and fleet planning is global; schema lacks reserved-space and vehicle-agency direct planning fields.", "Admin, Warehouse, Fleet owners", "Add reservation model and agency/vehicle planning linkage if required by operations."],
    ["P2", "Heatmap settings persistence", "Driver heatmap thresholds are service-memory/configuration based because there is no settings model for this slice.", "Driver, Admin", "Persist heatmap threshold settings and audit changes."],
    ["P2", "Multi-region rollout", "Country pricing/tax/cross-border code paths exist, but Neon multi-region deployment remains environment/provisioning work.", "Admin, Operations", "Provision staging regions and run inter-agency settlement/regression tests."],
    ["P2", "Documentation and test evidence", "QA docs exist, but each role needs current test evidence after latest auth, fleet, dispatch, and registration fixes.", "All users", "Create role smoke-test checklist and attach pass/fail evidence."],
]

MODULE_ROWS = [
    ["Identity & Access", "Implemented", "OTP-first login, email OTP + password continuation, registration approval, role redirects, session-bound JWTs, forced logout, suspicious login alerts.", "Keep auth UX regression tests current; verify provider failure messaging."],
    ["Booking & Dispatch", "Implemented + active extension", "Bookings, driver matching, driver assignment accept/reject, courier-company multi-stop movement, status transitions.", "Run lifecycle E2E through completion and payment/invoice triggers."],
    ["Fleet & Verification", "Implemented", "Owner-scoped fleet APIs, Kenya brand/model catalog, mandatory camera-image evidence, admin verification workflow.", "Live-camera enforcement is strongest on mobile; web capture should be device-tested."],
    ["Warehouse & Inventory", "Implemented base + pending advanced audit", "Warehouse dashboard, bins, inventory, scan, waybill, loss detection, listings and online bookings.", "Advanced PRD operations need line-by-line audit: dock slots, GRN, pick-pack, VAS, cold chain."],
    ["Finance", "Implemented + pending UX", "Invoices, billing, platform fees, reconciliation, wallets/payment hooks, rate cards, contracts.", "Revenue add-on UIs and provider sandbox testing remain the key gaps."],
    ["Operations & Admin", "Implemented", "Admin dashboards, analytics, audit, fraud, alerts, system health, workforce, warehouse listing review.", "High-volume data, permission isolation, and live monitoring tests remain."],
    ["Marketplace", "Implemented base + pending premium UX", "Partner onboarding, opportunities, bids/negotiation, commission tracking.", "Featured listing monetization UI is pending."],
    ["Offline / Mobile", "Partially implemented", "Offline scan queue, cached maps, mobile auth and role surfaces.", "Full mobile parity and device QA remain pending."],
]


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = p.add_run("Zito PRD Pending Features Report")
    title.bold = True
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor.from_string(DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = p.add_run("Role-by-role feature coverage, pending work, and testing focus")
    sub.font.size = Pt(12)
    sub.font.color.rgb = RGBColor.from_string(ACCENT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = p.add_run(f"Prepared from repo PRD tracker, user guides, testing guide, and current route map | {datetime.now():%d %B %Y}")
    meta.font.size = Pt(9)
    meta.font.color.rgb = RGBColor(90, 99, 118)

    add_callout(
        doc,
        "Executive Summary",
        "Core PRD coverage is implemented through the Phase 1 to Phase 5 slices in the repo tracker. The main remaining work is not a single missing phase; it is product hardening: revenue-module frontend screens, provider-backed E2E testing, mobile parity, advanced warehouse operations audit, and commercial-program ledger decisions.",
        LIGHT,
    )

    doc.add_heading("1. Priority Pending Backlog", level=1)
    add_table(
        doc,
        ["Priority", "Pending Area", "What Is Pending", "Users Impacted", "Recommended Next Step"],
        PRIORITY_ROWS,
        [0.55, 1.25, 2.55, 1.35, 2.05],
    )

    doc.add_heading("2. Pending Work by User / Role", level=1)
    add_table(
        doc,
        ["User / Role", "Current PRD Feature Coverage", "Pending / Refinement Needed", "Testing Focus"],
        ROLE_ROWS,
        [1.05, 2.45, 2.45, 2.15],
    )

    doc.add_heading("3. Module-Level Status", level=1)
    add_table(
        doc,
        ["PRD Module", "Status", "Current Coverage", "Pending Detail"],
        MODULE_ROWS,
        [1.35, 1.05, 3.0, 2.25],
    )

    doc.add_heading("4. Detailed User Feature Checklist", level=1)
    details = {
        "Customer": [
            "Register/login with OTP, create bookings, view booking history and booking detail.",
            "Track active shipments with cached map fallback and route visibility.",
            "Payments, invoices, support ticket creation and conversation detail.",
            "Owned fleet registration with mandatory camera evidence and admin verification.",
            "Warehouse discovery and online warehouse booking against approved listings.",
        ],
        "Corporate": [
            "Corporate booking and contract-aware commercial operation.",
            "Invoice and contract pages for billing/credit visibility.",
            "Owned fleet management under corporate account scope.",
            "Pending focus: contract exposure, credit-limit edges, consolidated billing regression.",
        ],
        "Driver": [
            "Dashboard, jobs, shift, heatmap, earnings, driver assignment accept/reject.",
            "Trip lifecycle updates and delivery verification paths.",
            "Pending focus: subscription tier UI, mobile identity verification photo QA, earnings reconciliation.",
        ],
        "Transporter / Agent": [
            "Fleet onboarding, driver linking, marketplace participation, invoices.",
            "Vehicle creation requires all images/documents and Kenya brand/model selection.",
            "Pending focus: featured-listing purchase UI, agent commission visibility, supply network isolation.",
        ],
        "Courier Company": [
            "Dispatch, bookings, new movement planning, scan ops, waybills, owned fleet, invoices.",
            "Capacity source supports Owned Fleet, CFA Network, and Blended execution.",
            "Pending focus: realistic parcel-volume chain-of-custody and mobile/web parity.",
        ],
        "Warehouse Partner": [
            "Dashboard, bins, inventory, scan, loss detection, listings, bookings, marketplace.",
            "Admin-reviewed public warehouse listings and customer-facing booking flow are active.",
            "Pending focus: advanced PRD warehouse operations, capacity reservation stress tests, mobile parity.",
        ],
        "Admin / Super Admin / Agency Staff": [
            "Admin controls users, fleet, KYC/fleet verification, marketplace, billing, finance, fraud, analytics, system health, workforce and warehouse listings.",
            "Agency staff handles agency-scoped accounts, operations and support routes.",
            "Pending focus: RBAC proof, department-specific permissions, audit evidence and revenue frontend completion.",
        ],
    }
    for role, items in details.items():
        doc.add_heading(role, level=2)
        add_bullets(doc, items)

    doc.add_heading("5. Recommended Execution Order", level=1)
    add_bullets(
        doc,
        [
            "Finish frontend screens for backend-ready revenue modules: driver subscriptions, featured listings, verification expedite.",
            "Run full role smoke tests using the seeded QA login matrix and record screenshots/results.",
            "Run provider-backed staging E2E for OTP, payment, SMS, uploads, maps, PDFs, scan/offline sync.",
            "Audit advanced warehouse PRD items against backend models/routes before adding more UI.",
            "Decide whether promo, loyalty and referral need dedicated ledger tables before scale.",
            "Complete mobile parity and device QA for Android/iPhone across customer, partner, internal and agency surfaces.",
        ],
    )

    add_callout(
        doc,
        "Important Scope Note",
        "This report reflects the current repo tracker and local code routes, not a production operations sign-off. Items marked implemented still require target-environment testing before release acceptance.",
        WARN,
    )

    doc.add_heading("6. Source References Used", level=1)
    add_bullets(
        doc,
        [
            "backend/PRD_TRACKER.md",
            "docs/prd/ZITO_PRD_v10_ULTIMATE.txt",
            "docs/qa/ZITO_USER_GUIDES.md",
            "docs/qa/TESTING_DOCUMENT.md",
            "frontend/src/app route map",
        ],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Zito PRD Pending Features Report | Aurenza Limited")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(110, 118, 135)

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_doc()
