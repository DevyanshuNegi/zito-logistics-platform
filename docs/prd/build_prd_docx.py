from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_TXT = ROOT / "ZITO_PRD_v10_ULTIMATE.txt"
OUTPUT_DOCX = ROOT / "ZITO_PRD_v10_ULTIMATE.docx"

NAVY = RGBColor(0x1B, 0x3F, 0x72)
AMBER = RGBColor(0xE8, 0xA0, 0x20)
TEXT = RGBColor(0x1A, 0x1A, 0x2E)
SUBTEXT = RGBColor(0x47, 0x55, 0x69)
BORDER = "C7D4E8"

HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)\.\s+(.*)$")
EMOJI_PREFIX_RE = re.compile(r"^[^\w(]+(?=\s)", re.UNICODE)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def remove_emoji_prefix(text: str) -> str:
    return EMOJI_PREFIX_RE.sub("", text).strip()


def is_label(line: str, next_line: str | None) -> bool:
    if not line or len(line) > 60:
        return False
    if line.startswith("- "):
        return False
    if HEADING_RE.match(line):
        return False
    if ":" in line:
        return False
    if next_line is None or not next_line.strip():
        return False
    if HEADING_RE.match(next_line) or next_line.startswith("- "):
        return False
    if line.isupper():
        return False
    return True


def apply_run_font(run, size: int, *, bold: bool = False, color: RGBColor = TEXT) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:cs"), "Arial")


def style_paragraph(paragraph, *, after: int = 4, before: int = 0, line: float = 1.15) -> None:
    pf = paragraph.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = line


def add_styled_paragraph(document: Document, text: str, *, size: int = 10, bold: bool = False,
                         color: RGBColor = TEXT, align=WD_ALIGN_PARAGRAPH.LEFT,
                         after: int = 4, before: int = 0, style_name: str | None = None):
    paragraph = document.add_paragraph(style=style_name)
    paragraph.alignment = align
    style_paragraph(paragraph, after=after, before=before)
    run = paragraph.add_run(text)
    apply_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.different_first_page_header_footer = True

    default_header = section.header
    p = default_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, after=2)
    run = p.add_run("ZITO PRD V10.3 - Ultimate Master")
    apply_run_font(run, size=8, bold=True, color=NAVY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(fp, after=0)
    run = fp.add_run("Confidential | ZITO Logistics Technologies | Synced from master PRD text")
    apply_run_font(run, size=8, color=SUBTEXT)
    return doc


def add_cover(doc: Document, lines: list[str]) -> None:
    add_styled_paragraph(doc, lines[0], size=30, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, before=90)
    add_styled_paragraph(doc, f"{lines[1]}  {lines[2]}", size=18, bold=True, color=TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_styled_paragraph(doc, lines[4], size=11, bold=True, color=AMBER, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = "Table Grid"
    meta_table.autofit = False
    widths = (Inches(2.0), Inches(4.8))
    rows = [
        ("Version", f"{lines[3]} | Integrated Ref: {lines[7].split(':', 1)[1].strip()}"),
        ("Classification", lines[5].split(":", 1)[1].strip()),
        ("Addenda", lines[6].split(":", 1)[1].strip()),
    ]
    for row_idx, (label, value) in enumerate(rows):
        left = meta_table.rows[row_idx].cells[0]
        right = meta_table.rows[row_idx].cells[1]
        left.width = widths[0]
        right.width = widths[1]
        set_cell_shading(left, "EEF4FF")
        for cell in (left, right):
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        left_run = left.paragraphs[0].add_run(label)
        apply_run_font(left_run, size=10, bold=True, color=NAVY)
        right_run = right.paragraphs[0].add_run(value)
        apply_run_font(right_run, size=10, color=TEXT)

    add_styled_paragraph(doc, "Copyright 2024 ZITO Logistics Technologies. All Rights Reserved.", size=9,
                         color=SUBTEXT, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=4)
    add_styled_paragraph(doc, "Master PRD export generated from repository text source on 2026-05-04.",
                         size=8, color=SUBTEXT, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc: Document, level: int, title: str) -> None:
    style_name = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
    }.get(level, "Heading 4")
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    apply_run_font(
        run,
        size={1: 15, 2: 12, 3: 11}.get(level, 10),
        bold=True,
        color=NAVY if level < 3 else SUBTEXT,
    )


def build_document() -> None:
    lines = SOURCE_TXT.read_text(encoding="utf-8").splitlines()
    doc = configure_document()
    add_cover(doc, lines[:8])

    start_idx = next(i for i, line in enumerate(lines) if HEADING_RE.match(line))
    content = lines[start_idx:]

    for idx, raw_line in enumerate(content):
        line = raw_line.strip()
        if not line:
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            number, title = heading_match.groups()
            level = min(len(number.split(".")), 4)
            clean_title = f"{number}. {remove_emoji_prefix(title)}"
            add_heading(doc, level, clean_title)
            continue

        next_line = None
        if idx + 1 < len(content):
            next_line = content[idx + 1].strip()

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(line[2:].strip())
            apply_run_font(run, size=9, color=TEXT)
            continue

        if is_label(line, next_line):
            add_styled_paragraph(doc, remove_emoji_prefix(line), size=10, bold=True, color=SUBTEXT, after=1, before=6)
            continue

        add_styled_paragraph(doc, line, size=9, color=TEXT, after=3)

    doc.core_properties.title = "ZITO PRD v10 Ultimate"
    doc.core_properties.subject = "Master product requirements document"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.company = "ZITO Logistics Technologies"
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
